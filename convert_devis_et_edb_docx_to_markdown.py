#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Conversion DOC/DOCX -> Markdown (dataset)
- NDC (colonne G) + EDB (colonne F)
- Filtre Excel : B=1, C=1, D=OUI, E=NON
- Ignore images
- Ignore page de garde / synthèse / tables des matières
- Conserve titres, paragraphes, sauts de lignes, listes multi-niveaux
- Conserve tableaux en Markdown ou HTML (supporte rowspan/colspan)
- Déduplication conservative (doublons consécutifs)
- Rapports CSV + logs d'erreurs

Dépendances:
  pip install -U pandas openpyxl python-docx lxml
Optionnel (pour lire .doc):
  pip install pywin32 (si MS Word installé)
  ou avoir LibreOffice accessible via 'soffice' dans le PATH
"""

from __future__ import annotations

import re
import os
import sys
import html
import shutil
import subprocess
import traceback
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple, Union, Set
from collections import defaultdict
from enum import Enum, auto

import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree

# Configuration du logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ------------------------------
# Configuration
# ------------------------------
EXCEL_NAME = "couverture_EDB_NDC_par_RITM.xlsx"

# Colonnes Excel par lettre -> index (A=0)
COL_B = 1
COL_C = 2
COL_D = 3
COL_E = 4
COL_F = 5  # EDB
COL_G = 6  # NDC

FILTER_B_EQ = 1
FILTER_C_EQ = 1
FILTER_D_EQ = "OUI"
FILTER_E_EQ = "NON"

OUTPUT_DIRNAME = "dataset_markdown"
LOG_DIRNAME = "_logs"
SUBDIR_NDC = "ndc"
SUBDIR_EDB = "edb"

# Namespace XML Word
WORD_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}

# Patterns pour identifier les sections à ignorer
TOC_PATTERNS = [
    r"^table\s+des\s+mati[eè]res?\s*$",
    r"^sommaire\s*$",
    r"^table\s+of\s+contents?\s*$",
]

SKIP_SECTION_PATTERNS = [
    r"^synth[eè]se\s*$",
    r"^r[eé]sum[eé]\s*$",
    r"^abstract\s*$",
    r"^executive\s+summary\s*$",
    r"^page\s+de\s+garde\s*$",
    r"^cover\s+page\s*$",
]

# Patterns pour les lignes de TOC (à ignorer)
TOC_LINE_PATTERNS = [
    # Ligne avec numéro de page à la fin et points de suite
    r"^.+?\.{2,}\s*\d+\s*$",
    # Ligne avec tabulation et numéro de page
    r"^.+?\t+\d+\s*$",
    # Numérotation romaine avec titre et numéro de page
    r"^\s*[IVXLCDM]+\.?\s+.+?\s+\d+\s*$",
    # Numérotation décimale avec titre et numéro de page
    r"^\s*\d+(?:\.\d+)*\.?\s+.+?\s+\d+\s*$",
]

# Patterns pour les titres (numérotation)
HEADING_PATTERNS = [
    # Numérotation romaine: "I. Titre" ou "I.1. Titre"
    (r"^\s*([IVXLCDM]+)\.\s+(.+)$", 1),
    (r"^\s*([IVXLCDM]+)\.(\d+)\.\s+(.+)$", 2),
    (r"^\s*([IVXLCDM]+)\.(\d+)\.(\d+)\.\s+(.+)$", 3),
    # Numérotation décimale: "1. Titre" ou "1.1 Titre" ou "1.1.1 Titre"
    (r"^\s*(\d+)\.\s+(.+)$", 1),
    (r"^\s*(\d+)\.(\d+)\.?\s+(.+)$", 2),
    (r"^\s*(\d+)\.(\d+)\.(\d+)\.?\s+(.+)$", 3),
    (r"^\s*(\d+)\.(\d+)\.(\d+)\.(\d+)\.?\s+(.+)$", 4),
]

# Titres connus (normalisés en minuscules)
KNOWN_SECTION_TITLES = {
    # Niveau 1 - Sections principales NDC
    "description du projet": 1,
    "périmètre du projet": 1,
    "perimetre du projet": 1,
    "securité et réglementation": 1,
    "sécurité et réglementation": 1,
    "contraintes et risques": 1,
    "description technique de la solution": 1,
    "démarche pour la mise en œuvre": 1,
    "demarche pour la mise en oeuvre": 1,
    "offre de service en fonctionnement": 1,
    "evaluation financière": 1,
    "évaluation financière": 1,
    "gestion de la documentation du projet": 1,
    "rse - impact co2": 1,
    "rse – impact co2": 1,
    "introduction": 1,
    "conclusion": 1,
    "annexes": 1,

    # Niveau 2 - Sous-sections
    "contexte": 2,
    "le besoin exprimé": 2,
    "besoin exprimé": 2,
    "objectifs du projet": 2,
    "objectifs": 2,
    "périmètre": 2,
    "perimetre": 2,
    "hors périmètre": 2,
    "hors perimetre": 2,
    "projet partenaire": 2,
    "projet interne": 2,
    "contrôle psee": 2,
    "controle psee": 2,
    "contraintes et prérequis": 2,
    "contraintes et prerequis": 2,
    "contraintes": 2,
    "prérequis": 2,
    "prerequis": 2,
    "risques projet": 2,
    "risques": 2,
    "description de la solution": 2,
    "architecture": 2,
    "composants et dimensionnement": 2,
    "dimensionnement": 2,
    "lotissement": 2,
    "livrables du projet": 2,
    "livrables": 2,
    "jalons clés du projet": 2,
    "jalons cles du projet": 2,
    "jalons": 2,
    "macro-planning": 2,
    "macro planning": 2,
    "planning": 2,
    "détails des contributions": 2,
    "details des contributions": 2,
    "contributions": 2,
    "comitologie": 2,
    "validation de la solution": 2,
    "niveaux de service": 2,
    "coûts du projet": 2,
    "couts du projet": 2,
    "coûts": 2,
    "couts": 2,
    "facturation": 2,
    "coûts de fonctionnement": 2,
    "couts de fonctionnement": 2,
    "coûts de construction": 2,
    "couts de construction": 2,
}


# ------------------------------
# Classes utilitaires
# ------------------------------
class DocState(Enum):
    """État de la machine d'état pour le parsing du document."""
    PREAMBLE = auto()      # Page de garde, avant-propos
    TOC = auto()           # Table des matières
    SKIP_SECTION = auto()  # Section à ignorer (synthèse, etc.)
    BODY = auto()          # Contenu principal


@dataclass
class ParsedBlock:
    """Représente un bloc parsé du document."""
    type: str  # 'heading', 'paragraph', 'list_item', 'table', 'empty'
    content: str
    level: int = 0  # Pour les headings et les listes
    list_type: str = ""  # 'bullet' ou 'ordered'
    raw_text: str = ""  # Texte brut sans formatage


@dataclass
class NumberingInfo:
    """Information sur la numérotation d'un paragraphe."""
    num_id: int
    ilvl: int
    num_format: str  # 'bullet', 'decimal', 'lowerLetter', etc.


# ------------------------------
# Fonctions utilitaires XML
# ------------------------------
def xpath(element, expr: str) -> list:
    """Execute XPath sur un élément lxml avec les namespaces Word."""
    try:
        return element.xpath(expr, namespaces=WORD_NS)
    except Exception:
        return []


def get_text_recursive(element) -> str:
    """Récupère tout le texte d'un élément XML récursivement."""
    texts = []
    for node in element.iter():
        if node.text:
            texts.append(node.text)
        if node.tail:
            texts.append(node.tail)
    return "".join(texts)


# ------------------------------
# Conversion DOC -> DOCX
# ------------------------------
def convert_doc_to_docx(input_doc: Path, workdir: Path) -> Optional[Path]:
    """
    Convertit un fichier .doc en .docx.
    Utilise MS Word via COM (Windows) ou LibreOffice.
    """
    workdir.mkdir(parents=True, exist_ok=True)
    out_docx = workdir / (input_doc.stem + ".docx")

    if out_docx.exists():
        return out_docx

    # Essayer MS Word via COM (Windows)
    if os.name == "nt":
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(str(input_doc.resolve()))
            doc.SaveAs(str(out_docx.resolve()), FileFormat=16)
            doc.Close(False)
            word.Quit()
            if out_docx.exists():
                logger.info(f"Converti via MS Word: {input_doc.name}")
                return out_docx
        except Exception as e:
            logger.debug(f"Conversion MS Word échouée: {e}")

    # Essayer LibreOffice
    soffice = shutil.which("soffice") or shutil.which("soffice.exe")
    if soffice:
        try:
            cmd = [
                soffice, "--headless", "--nologo",
                "--convert-to", "docx",
                "--outdir", str(workdir),
                str(input_doc)
            ]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            if out_docx.exists():
                logger.info(f"Converti via LibreOffice: {input_doc.name}")
                return out_docx
        except Exception as e:
            logger.debug(f"Conversion LibreOffice échouée: {e}")

    return None


# ------------------------------
# Extraction des informations de numérotation
# ------------------------------
class NumberingExtractor:
    """Extrait les informations de numérotation du document Word."""

    def __init__(self, doc: Document):
        self.doc = doc
        self.formats: Dict[Tuple[int, int], str] = {}
        self._extract_numbering()

    def _extract_numbering(self):
        """Parse le fichier numbering.xml pour extraire les formats."""
        try:
            numbering_part = self.doc.part.numbering_part
            if numbering_part is None:
                return

            numbering_xml = numbering_part.element

            # Mapper abstractNumId -> formats par niveau
            abstract_formats: Dict[str, Dict[str, str]] = {}

            for abstract_num in xpath(numbering_xml, ".//w:abstractNum"):
                abs_id = abstract_num.get(qn("w:abstractNumId"))
                if not abs_id:
                    continue

                level_formats = {}
                for lvl in xpath(abstract_num, "./w:lvl"):
                    ilvl = lvl.get(qn("w:ilvl"))
                    num_fmt_elem = xpath(lvl, "./w:numFmt/@w:val")
                    if num_fmt_elem:
                        level_formats[ilvl] = num_fmt_elem[0]

                abstract_formats[abs_id] = level_formats

            # Mapper numId -> abstractNumId
            num_to_abstract: Dict[str, str] = {}
            for num in xpath(numbering_xml, ".//w:num"):
                num_id = num.get(qn("w:numId"))
                abs_id_refs = xpath(num, "./w:abstractNumId/@w:val")
                if num_id and abs_id_refs:
                    num_to_abstract[num_id] = abs_id_refs[0]

            # Construire le mapping final
            for num_id, abs_id in num_to_abstract.items():
                if abs_id in abstract_formats:
                    for ilvl, fmt in abstract_formats[abs_id].items():
                        self.formats[(int(num_id), int(ilvl))] = fmt

        except Exception as e:
            logger.debug(f"Erreur extraction numérotation: {e}")

    def get_list_info(self, paragraph: Paragraph) -> Optional[NumberingInfo]:
        """Retourne les infos de liste pour un paragraphe."""
        p_elem = paragraph._p
        pPr = p_elem.pPr
        if pPr is None:
            return None

        numPr = pPr.numPr
        if numPr is None:
            # Vérifier le style
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            if any(kw in style_name for kw in ["list", "puce", "bullet", "enum", "num"]):
                is_bullet = any(kw in style_name for kw in ["puce", "bullet"])
                return NumberingInfo(0, 0, "bullet" if is_bullet else "decimal")
            return None

        num_id_elem = xpath(numPr, "./w:numId/@w:val")
        ilvl_elem = xpath(numPr, "./w:ilvl/@w:val")

        if not num_id_elem:
            return None

        num_id = int(num_id_elem[0])
        ilvl = int(ilvl_elem[0]) if ilvl_elem else 0

        fmt = self.formats.get((num_id, ilvl), "decimal")

        return NumberingInfo(num_id, ilvl, fmt)


# ------------------------------
# Extraction du texte des paragraphes
# ------------------------------
class TextExtractor:
    """Extrait le texte formaté des paragraphes Word."""

    @staticmethod
    def has_drawing(run_element) -> bool:
        """Vérifie si un run contient une image/dessin."""
        return bool(
            xpath(run_element, ".//w:drawing") or
            xpath(run_element, ".//w:pict") or
            xpath(run_element, ".//pic:pic")
        )

    @staticmethod
    def extract_paragraph_text(paragraph: Paragraph, include_formatting: bool = True) -> str:
        """
        Extrait le texte d'un paragraphe avec formatage optionnel (bold, italic).
        """
        p_elem = paragraph._p
        chunks: List[str] = []

        for run in xpath(p_elem, "./w:r"):
            # Ignorer les runs avec images
            if TextExtractor.has_drawing(run):
                continue

            # Récupérer le texte
            text_elements = xpath(run, "./w:t")
            text = "".join([(t.text or "") for t in text_elements])

            if not text:
                # Vérifier les sauts de ligne
                if xpath(run, "./w:br"):
                    chunks.append("\n")
                continue

            # Nettoyer le texte
            text = text.replace("\u00A0", " ")  # Non-breaking space
            text = text.replace("\u200B", "")   # Zero-width space

            if include_formatting:
                # Vérifier le formatage
                rPr = xpath(run, "./w:rPr")
                is_bold = bool(xpath(run, "./w:rPr/w:b[not(@w:val='false')]"))
                is_italic = bool(xpath(run, "./w:rPr/w:i[not(@w:val='false')]"))

                if is_bold and is_italic:
                    text = f"***{text}***"
                elif is_bold:
                    text = f"**{text}**"
                elif is_italic:
                    text = f"*{text}*"

            chunks.append(text)

        result = "".join(chunks).strip()

        # Nettoyer les formatages mal fermés/ouverts
        result = re.sub(r'\*{4,}', '**', result)

        return result

    @staticmethod
    def extract_raw_text(paragraph: Paragraph) -> str:
        """Extrait le texte brut sans formatage."""
        return TextExtractor.extract_paragraph_text(paragraph, include_formatting=False)


# ------------------------------
# Détection des titres
# ------------------------------
class HeadingDetector:
    """Détecte et classifie les titres."""

    @staticmethod
    def get_style_heading_level(paragraph: Paragraph) -> Optional[int]:
        """Retourne le niveau de titre basé sur le style Word."""
        if not paragraph.style:
            return None

        style_name = paragraph.style.name or ""
        style_lower = style_name.lower()

        # Styles de titre standards: "Heading 1", "Titre 1", etc.
        match = re.match(r"^(heading|titre|title)\s*(\d+)", style_lower)
        if match:
            level = int(match.group(2))
            return min(6, max(1, level))

        # Vérifier outlineLevel dans les propriétés du paragraphe
        p_elem = paragraph._p
        pPr = p_elem.pPr
        if pPr is not None:
            outline_lvl = xpath(pPr, "./w:outlineLvl/@w:val")
            if outline_lvl:
                level = int(outline_lvl[0]) + 1
                return min(6, max(1, level))

        return None

    @staticmethod
    def get_pattern_heading_level(text: str) -> Optional[Tuple[int, str]]:
        """
        Détecte un titre par pattern de numérotation.
        Retourne (niveau, texte nettoyé) ou None.
        """
        text = text.strip()
        if not text:
            return None

        for pattern, base_level in HEADING_PATTERNS:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                groups = match.groups()
                # Le titre est le dernier groupe
                title = groups[-1].strip()
                # Ne pas considérer comme titre si le texte est trop long
                if len(title) > 200:
                    continue
                # Vérifier que ce n'est pas une ligne de TOC (numéro de page à la fin)
                if re.search(r'\s+\d+\s*$', title) and len(title) < 100:
                    continue
                return (min(6, base_level), title)

        return None

    @staticmethod
    def get_known_title_level(text: str) -> Optional[int]:
        """Retourne le niveau si c'est un titre connu."""
        normalized = text.strip().lower()
        # Nettoyer la ponctuation
        normalized = re.sub(r'[:\-–—]+$', '', normalized).strip()
        return KNOWN_SECTION_TITLES.get(normalized)

    @staticmethod
    def detect_heading(paragraph: Paragraph, text: str) -> Optional[Tuple[int, str]]:
        """
        Détecte si un paragraphe est un titre et retourne (niveau, texte).
        """
        if not text or not text.strip():
            return None

        text = text.strip()

        # 1. Vérifier le style Word
        style_level = HeadingDetector.get_style_heading_level(paragraph)
        if style_level is not None:
            return (style_level, text)

        # 2. Vérifier les patterns de numérotation
        pattern_result = HeadingDetector.get_pattern_heading_level(text)
        if pattern_result:
            return pattern_result

        # 3. Vérifier les titres connus
        known_level = HeadingDetector.get_known_title_level(text)
        if known_level is not None:
            return (known_level, text)

        return None


# ------------------------------
# Détection des sections à ignorer
# ------------------------------
class SectionFilter:
    """Filtre les sections à ignorer (TOC, préambule, etc.)."""

    @staticmethod
    def is_toc_heading(text: str) -> bool:
        """Vérifie si le texte est un titre de table des matières."""
        normalized = text.strip().lower()
        for pattern in TOC_PATTERNS:
            if re.match(pattern, normalized):
                return True
        return False

    @staticmethod
    def is_skip_section_heading(text: str) -> bool:
        """Vérifie si c'est une section à ignorer (synthèse, etc.)."""
        normalized = text.strip().lower()
        for pattern in SKIP_SECTION_PATTERNS:
            if re.match(pattern, normalized):
                return True
        return False

    @staticmethod
    def is_toc_line(text: str) -> bool:
        """Vérifie si une ligne ressemble à une entrée de TOC."""
        if not text or not text.strip():
            return False

        text = text.strip()

        # Ligne avec tabulation (typique des TOC)
        if "\t" in text:
            return True

        # Points de suite
        if "..." in text or "…" in text:
            return True

        for pattern in TOC_LINE_PATTERNS:
            if re.match(pattern, text, re.IGNORECASE):
                return True

        return False

    @staticmethod
    def is_page_number_only(text: str) -> bool:
        """Vérifie si le texte est juste un numéro de page."""
        return bool(re.match(r'^\s*\d+\s*$', text.strip()))

    @staticmethod
    def looks_like_cover_page_element(text: str) -> bool:
        """Vérifie si ça ressemble à un élément de page de garde."""
        text_lower = text.strip().lower()

        cover_patterns = [
            r"^confidentiel",
            r"^document\s+interne",
            r"^usage\s+interne",
            r"^version\s*:?\s*[\d\.]+",
            r"^date\s*:?\s*\d",
            r"^auteur\s*:?\s*",
            r"^rédacteur\s*:?\s*",
            r"^redacteur\s*:?\s*",
            r"^statut\s*:?\s*",
            r"^référence\s*:?\s*",
            r"^reference\s*:?\s*",
            r"^diffusion\s*:?\s*",
            r"^\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}$",  # Date seule
        ]

        for pattern in cover_patterns:
            if re.match(pattern, text_lower):
                return True

        return False


# ------------------------------
# Conversion des tableaux
# ------------------------------
class TableConverter:
    """Convertit les tableaux Word en Markdown/HTML."""

    @staticmethod
    def extract_cell_text(tc_element) -> str:
        """Extrait le texte d'une cellule de tableau."""
        paragraphs = xpath(tc_element, ".//w:p")
        parts = []

        for p in paragraphs:
            text_parts = []
            for run in xpath(p, ".//w:r"):
                # Ignorer les images
                if xpath(run, ".//w:drawing") or xpath(run, ".//w:pict"):
                    continue
                texts = xpath(run, ".//w:t")
                text_parts.append("".join([(t.text or "") for t in texts]))

            para_text = "".join(text_parts).strip()
            if para_text:
                parts.append(para_text)

        return " | ".join(parts).replace("\u00A0", " ").strip()

    @staticmethod
    def get_cell_span(tc_element) -> Tuple[int, int, bool]:
        """
        Retourne (colspan, rowspan_start, is_merged_continue).
        """
        # Colspan (gridSpan)
        grid_span = xpath(tc_element, "./w:tcPr/w:gridSpan/@w:val")
        colspan = int(grid_span[0]) if grid_span else 1

        # Rowspan (vMerge)
        vmerge = xpath(tc_element, "./w:tcPr/w:vMerge")
        if vmerge:
            vmerge_val = xpath(tc_element, "./w:tcPr/w:vMerge/@w:val")
            if vmerge_val and vmerge_val[0] == "restart":
                return (colspan, True, False)  # Début d'un rowspan
            else:
                return (colspan, False, True)  # Continuation d'un rowspan

        return (colspan, False, False)

    @staticmethod
    def table_to_markdown(table: Table) -> str:
        """Convertit un tableau en Markdown simple."""
        tbl_elem = table._tbl
        rows = xpath(tbl_elem, "./w:tr")

        if not rows:
            return ""

        md_rows = []
        max_cols = 0

        for row in rows:
            cells = xpath(row, "./w:tc")
            row_texts = []
            col_count = 0

            for cell in cells:
                colspan, _, is_continue = TableConverter.get_cell_span(cell)

                if is_continue:
                    # Cellule fusionnée verticalement - mettre du texte vide
                    row_texts.append("")
                else:
                    text = TableConverter.extract_cell_text(cell)
                    # Nettoyer le texte pour Markdown
                    text = text.replace("|", "\\|").replace("\n", " ")
                    row_texts.append(text)

                # Ajouter des cellules vides pour le colspan
                for _ in range(colspan - 1):
                    row_texts.append("")

                col_count += colspan

            max_cols = max(max_cols, col_count)
            md_rows.append(row_texts)

        # Normaliser le nombre de colonnes
        for row in md_rows:
            while len(row) < max_cols:
                row.append("")

        if not md_rows:
            return ""

        # Construire le Markdown
        lines = []

        # Première ligne (header)
        header = "| " + " | ".join(md_rows[0]) + " |"
        lines.append(header)

        # Séparateur
        sep = "| " + " | ".join(["---"] * max_cols) + " |"
        lines.append(sep)

        # Autres lignes
        for row in md_rows[1:]:
            line = "| " + " | ".join(row) + " |"
            lines.append(line)

        return "\n".join(lines)

    @staticmethod
    def table_to_html(table: Table) -> str:
        """Convertit un tableau en HTML (pour les tableaux complexes avec rowspan/colspan)."""
        tbl_elem = table._tbl
        rows = xpath(tbl_elem, "./w:tr")

        if not rows:
            return ""

        # Analyser la structure du tableau
        grid: List[List[Optional[dict]]] = []
        max_cols = 0

        for row_idx, row in enumerate(rows):
            cells = xpath(row, "./w:tc")
            row_data = []
            col_idx = 0

            for cell in cells:
                colspan, is_rowspan_start, is_rowspan_continue = TableConverter.get_cell_span(cell)
                text = TableConverter.extract_cell_text(cell)

                cell_info = {
                    "text": html.escape(text),
                    "colspan": colspan,
                    "rowspan": 1,
                    "is_continue": is_rowspan_continue,
                    "is_start": is_rowspan_start,
                }

                for _ in range(colspan):
                    row_data.append(cell_info if _ == 0 else {"span_ref": cell_info})
                    col_idx += 1

            max_cols = max(max_cols, col_idx)
            grid.append(row_data)

        # Calculer les rowspans
        for col in range(max_cols):
            current_start = None
            for row_idx in range(len(grid)):
                if col >= len(grid[row_idx]):
                    continue
                cell = grid[row_idx][col]
                if cell is None:
                    continue
                if "span_ref" in cell:
                    continue
                if cell.get("is_start"):
                    current_start = (row_idx, col)
                elif cell.get("is_continue") and current_start:
                    start_row, start_col = current_start
                    if start_col < len(grid[start_row]):
                        start_cell = grid[start_row][start_col]
                        if "span_ref" not in start_cell:
                            start_cell["rowspan"] += 1
                    cell["skip"] = True
                else:
                    current_start = None

        # Générer le HTML
        html_lines = ["<table>"]

        for row_idx, row_data in enumerate(grid):
            html_lines.append("  <tr>")
            col_idx = 0
            while col_idx < len(row_data):
                cell = row_data[col_idx]
                if cell is None or cell.get("skip") or "span_ref" in cell:
                    col_idx += 1
                    continue

                attrs = []
                if cell["colspan"] > 1:
                    attrs.append(f'colspan="{cell["colspan"]}"')
                if cell["rowspan"] > 1:
                    attrs.append(f'rowspan="{cell["rowspan"]}"')

                attr_str = (" " + " ".join(attrs)) if attrs else ""
                html_lines.append(f'    <td{attr_str}>{cell["text"]}</td>')
                col_idx += cell["colspan"]

            html_lines.append("  </tr>")

        html_lines.append("</table>")
        return "\n".join(html_lines)

    @staticmethod
    def has_complex_structure(table: Table) -> bool:
        """Vérifie si le tableau a des cellules fusionnées."""
        tbl_elem = table._tbl

        # Vérifier les vMerge (fusion verticale)
        if xpath(tbl_elem, ".//w:vMerge"):
            return True

        # Vérifier les gridSpan > 1 (fusion horizontale)
        for span in xpath(tbl_elem, ".//w:gridSpan/@w:val"):
            if int(span) > 1:
                return True

        return False

    @staticmethod
    def contains_toc(table: Table) -> bool:
        """Vérifie si le tableau contient une table des matières."""
        tbl_elem = table._tbl
        for tc in xpath(tbl_elem, ".//w:tc"):
            text = TableConverter.extract_cell_text(tc)
            if SectionFilter.is_toc_heading(text):
                return True
        return False


# ------------------------------
# Itération sur les blocs du document
# ------------------------------
def iter_document_blocks(doc: Document) -> Iterable[Union[Paragraph, Table]]:
    """Itère sur les paragraphes et tableaux dans l'ordre du document."""
    body = doc.element.body

    for child in body.iterchildren():
        tag = child.tag

        if tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif tag == qn("w:tbl"):
            yield Table(child, doc)


# ------------------------------
# Conversion principale
# ------------------------------
class DocxToMarkdownConverter:
    """Convertisseur DOCX vers Markdown."""

    def __init__(self, docx_path: Path, mode: str = "ndc"):
        self.docx_path = docx_path
        self.mode = mode
        self.doc = Document(str(docx_path))
        self.numbering = NumberingExtractor(self.doc)

        self.state = DocState.PREAMBLE
        self.md_lines: List[str] = []
        self.in_list = False
        self.prev_list_level = -1
        self.skip_until_next_h1 = False
        self.preamble_line_count = 0
        self.body_started = False

    def convert(self) -> str:
        """Effectue la conversion et retourne le Markdown."""
        for block in iter_document_blocks(self.doc):
            self._process_block(block)

        return self._finalize()

    def _process_block(self, block: Union[Paragraph, Table]):
        """Traite un bloc (paragraphe ou tableau)."""
        if isinstance(block, Table):
            self._process_table(block)
        else:
            self._process_paragraph(block)

    def _process_table(self, table: Table):
        """Traite un tableau."""
        # Vérifier si c'est une TOC dans un tableau
        if TableConverter.contains_toc(table):
            self.state = DocState.TOC
            return

        # Ignorer les tableaux avant le corps
        if self.state in (DocState.PREAMBLE, DocState.TOC, DocState.SKIP_SECTION):
            return

        # Fermer la liste en cours si nécessaire
        self._close_list()

        # Convertir le tableau
        if TableConverter.has_complex_structure(table):
            table_md = TableConverter.table_to_html(table)
        else:
            table_md = TableConverter.table_to_markdown(table)

        if table_md:
            self.md_lines.append("")
            self.md_lines.append(table_md)
            self.md_lines.append("")

    def _process_paragraph(self, paragraph: Paragraph):
        """Traite un paragraphe."""
        # Extraire le texte
        text = TextExtractor.extract_paragraph_text(paragraph)
        raw_text = TextExtractor.extract_raw_text(paragraph)

        # Paragraphe vide
        if not raw_text.strip():
            if self.state == DocState.BODY:
                self._close_list()
                self.md_lines.append("")
            return

        # Machine d'état
        if self.state == DocState.PREAMBLE:
            self._handle_preamble(paragraph, text, raw_text)
        elif self.state == DocState.TOC:
            self._handle_toc(paragraph, text, raw_text)
        elif self.state == DocState.SKIP_SECTION:
            self._handle_skip_section(paragraph, text, raw_text)
        else:  # BODY
            self._handle_body(paragraph, text, raw_text)

    def _handle_preamble(self, paragraph: Paragraph, text: str, raw_text: str):
        """Gère le préambule (page de garde, avant-propos)."""
        self.preamble_line_count += 1

        # Vérifier si c'est un titre de TOC
        if SectionFilter.is_toc_heading(raw_text):
            self.state = DocState.TOC
            return

        # Vérifier si c'est une section à ignorer
        if SectionFilter.is_skip_section_heading(raw_text):
            self.state = DocState.SKIP_SECTION
            return

        # Vérifier si c'est un vrai titre de section (début du corps)
        heading = HeadingDetector.detect_heading(paragraph, raw_text)
        if heading:
            level, title = heading
            # Si c'est un H1 ou H2 significatif, c'est le début du corps
            if level <= 2 and not SectionFilter.looks_like_cover_page_element(title):
                self.state = DocState.BODY
                self.body_started = True
                self._add_heading(level, text)
                return

        # Si on a passé beaucoup de lignes sans trouver de TOC,
        # et qu'on voit du contenu substantiel, passer au corps
        if self.preamble_line_count > 50:
            # Probablement pas de TOC explicite
            if len(raw_text) > 100 and not SectionFilter.looks_like_cover_page_element(raw_text):
                self.state = DocState.BODY
                self.body_started = True
                self._handle_body(paragraph, text, raw_text)

    def _handle_toc(self, paragraph: Paragraph, text: str, raw_text: str):
        """Gère la table des matières (à ignorer)."""
        # Vérifier si on est sorti de la TOC
        if SectionFilter.is_toc_line(raw_text):
            return  # Encore dans la TOC

        if SectionFilter.is_page_number_only(raw_text):
            return  # Numéro de page seul

        # Vérifier si c'est un titre
        heading = HeadingDetector.detect_heading(paragraph, raw_text)
        if heading:
            level, title = heading
            # Ne pas prendre "Sommaire" ou "Table des matières" comme fin de TOC
            if not SectionFilter.is_toc_heading(raw_text) and not SectionFilter.is_skip_section_heading(raw_text):
                self.state = DocState.BODY
                self.body_started = True
                self._add_heading(level, text)
                return

        # Si c'est du contenu substantiel, passer au corps
        if len(raw_text) > 50 and not SectionFilter.is_toc_line(raw_text):
            self.state = DocState.BODY
            self.body_started = True
            self._handle_body(paragraph, text, raw_text)

    def _handle_skip_section(self, paragraph: Paragraph, text: str, raw_text: str):
        """Gère une section à ignorer (synthèse, etc.)."""
        # Vérifier si c'est un nouveau titre de niveau 1 (sortie de la section)
        heading = HeadingDetector.detect_heading(paragraph, raw_text)
        if heading:
            level, title = heading
            if level == 1 and not SectionFilter.is_skip_section_heading(raw_text):
                self.state = DocState.BODY
                self.body_started = True
                self._add_heading(level, text)
                return

        # Vérifier si c'est un titre de TOC
        if SectionFilter.is_toc_heading(raw_text):
            self.state = DocState.TOC

    def _handle_body(self, paragraph: Paragraph, text: str, raw_text: str):
        """Gère le contenu principal du document."""
        # Ignorer les éléments de page de garde qui traînent
        if not self.body_started or self.preamble_line_count < 5:
            if SectionFilter.looks_like_cover_page_element(raw_text):
                return

        self.body_started = True

        # Vérifier si c'est une section à ignorer
        if SectionFilter.is_skip_section_heading(raw_text):
            self.state = DocState.SKIP_SECTION
            return

        # Vérifier si c'est un titre
        heading = HeadingDetector.detect_heading(paragraph, raw_text)
        if heading:
            self._close_list()
            level, title = heading
            self._add_heading(level, text)
            return

        # Vérifier si c'est une liste
        list_info = self.numbering.get_list_info(paragraph)
        if list_info:
            self._add_list_item(list_info, text)
            return

        # Paragraphe normal
        self._close_list()
        self.md_lines.append(text)
        self.md_lines.append("")

    def _add_heading(self, level: int, text: str):
        """Ajoute un titre."""
        self.md_lines.append("")
        prefix = "#" * level
        self.md_lines.append(f"{prefix} {text}")
        self.md_lines.append("")

    def _add_list_item(self, list_info: NumberingInfo, text: str):
        """Ajoute un élément de liste."""
        if not self.in_list:
            self.md_lines.append("")
            self.in_list = True

        # Indentation basée sur le niveau
        indent = "  " * list_info.ilvl

        # Marqueur
        if list_info.num_format == "bullet":
            marker = "-"
        else:
            marker = "1."

        self.md_lines.append(f"{indent}{marker} {text}")
        self.prev_list_level = list_info.ilvl

    def _close_list(self):
        """Ferme une liste en cours."""
        if self.in_list:
            self.md_lines.append("")
            self.in_list = False
            self.prev_list_level = -1

    def _finalize(self) -> str:
        """Finalise et nettoie le Markdown."""
        self._close_list()

        # Dédupliquer les lignes consécutives identiques
        deduped = []
        prev_key = None
        for line in self.md_lines:
            key = re.sub(r'\s+', ' ', line.strip())
            if key and key == prev_key:
                continue
            deduped.append(line)
            prev_key = key if key else prev_key

        # Joindre et nettoyer
        result = "\n".join(deduped)

        # Supprimer les lignes vides excessives (max 2 consécutives)
        result = re.sub(r'\n{4,}', '\n\n\n', result)

        # Supprimer les lignes vides au début
        result = result.lstrip('\n')

        # Assurer une ligne vide à la fin
        result = result.rstrip() + '\n'

        return result


# ------------------------------
# Chargement Excel
# ------------------------------
def load_targets_from_excel(excel_path: Path) -> Tuple[List[str], List[str]]:
    """Charge les fichiers à traiter depuis Excel."""
    df = pd.read_excel(excel_path, engine="openpyxl")

    b = df.iloc[:, COL_B]
    c = df.iloc[:, COL_C]
    d = df.iloc[:, COL_D]
    e = df.iloc[:, COL_E]
    f = df.iloc[:, COL_F]  # EDB
    g = df.iloc[:, COL_G]  # NDC

    d_norm = d.astype(str).str.strip().str.upper()
    e_norm = e.astype(str).str.strip().str.upper()

    mask = (b == FILTER_B_EQ) & (c == FILTER_C_EQ) & (d_norm == FILTER_D_EQ) & (e_norm == FILTER_E_EQ)

    edb = f[mask].dropna().astype(str).tolist()
    ndc = g[mask].dropna().astype(str).tolist()

    def clean(lst):
        out = []
        for t in lst:
            t = t.strip().strip('"').strip("'")
            if t:
                out.append(t)
        return out

    return clean(ndc), clean(edb)


# ------------------------------
# Programme principal
# ------------------------------
def main() -> int:
    cwd = Path(".").resolve()
    excel_path = cwd / EXCEL_NAME

    base_out = cwd / OUTPUT_DIRNAME
    log_dir = base_out / LOG_DIRNAME
    out_ndc = base_out / SUBDIR_NDC
    out_edb = base_out / SUBDIR_EDB
    tmp_conv = base_out / "_tmp_doc_conversion"

    # Créer les répertoires
    for d in [base_out, log_dir, out_ndc, out_edb, tmp_conv]:
        d.mkdir(parents=True, exist_ok=True)

    if not excel_path.exists():
        logger.error(f"Fichier Excel introuvable: {excel_path}")
        return 2

    ndc_list, edb_list = load_targets_from_excel(excel_path)

    logger.info(f"Fichiers NDC à traiter: {len(ndc_list)}")
    logger.info(f"Fichiers EDB à traiter: {len(edb_list)}")

    if not ndc_list and not edb_list:
        logger.info("Aucun fichier à traiter après filtrage Excel.")
        return 0

    report_rows = []
    stats = {"ok": 0, "error": 0, "missing": 0}

    def process_file(name: str, mode: str):
        """Traite un fichier."""
        src = cwd / name
        ext = src.suffix.lower()

        # Si pas d'extension, assumer .docx
        if not ext:
            src = src.with_suffix(".docx")
            ext = ".docx"

        out_dir = out_ndc if mode == "ndc" else out_edb

        # Chercher le fichier avec différentes extensions
        if not src.exists():
            for try_ext in [".docx", ".doc", ".DOCX", ".DOC"]:
                alt = src.with_suffix(try_ext)
                if alt.exists():
                    src = alt
                    ext = try_ext.lower()
                    break

        if not src.exists():
            stats["missing"] += 1
            report_rows.append((mode, name, str(src), "", "MISSING", "Fichier introuvable"))
            logger.warning(f"Fichier introuvable: {src}")
            return

        # Convertir .doc si nécessaire
        working_docx = src
        if ext == ".doc":
            converted = convert_doc_to_docx(src, tmp_conv)
            if not converted:
                stats["error"] += 1
                report_rows.append((mode, name, str(src), "", "CONVERT_ERROR",
                                   "Impossible de convertir .doc -> .docx"))
                logger.error(f"Conversion .doc échouée: {src.name}")
                return
            working_docx = converted

        try:
            # Convertir en Markdown
            converter = DocxToMarkdownConverter(working_docx, mode=mode)
            md_content = converter.convert()

            # Écrire le fichier
            out_name = Path(name).stem + ".md"
            out_path = out_dir / out_name
            out_path.write_text(md_content, encoding="utf-8")

            stats["ok"] += 1
            report_rows.append((mode, name, str(src), str(out_path), "OK", ""))
            logger.info(f"[OK] ({mode}) {src.name} -> {out_path.name}")

        except Exception as ex:
            stats["error"] += 1
            err_msg = f"{type(ex).__name__}: {ex}"
            report_rows.append((mode, name, str(src), "", "ERROR", err_msg))

            # Sauvegarder le traceback
            trace = traceback.format_exc()
            log_file = log_dir / (Path(name).stem + f".{mode}.error.log")
            log_file.write_text(trace, encoding="utf-8")

            logger.error(f"[ERREUR] ({mode}) {src.name}: {err_msg}")

    # Traiter les fichiers
    for f in ndc_list:
        process_file(f, mode="ndc")

    for f in edb_list:
        process_file(f, mode="edb")

    # Générer le rapport
    report_df = pd.DataFrame(
        report_rows,
        columns=["type", "source_excel", "input_path", "output_md", "status", "error"]
    )
    report_path = base_out / "conversion_report.csv"
    report_df.to_csv(report_path, index=False, encoding="utf-8")

    # Résumé
    logger.info("")
    logger.info("=== Résumé ===")
    logger.info(f"Traités avec succès: {stats['ok']}")
    logger.info(f"Fichiers manquants: {stats['missing']}")
    logger.info(f"Erreurs: {stats['error']}")
    logger.info(f"Sorties NDC: {out_ndc}")
    logger.info(f"Sorties EDB: {out_edb}")
    logger.info(f"Rapport: {report_path}")

    if stats["error"] > 0:
        logger.info(f"Logs d'erreurs: {log_dir}")

    return 0 if stats["error"] == 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())
