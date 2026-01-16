#!/usr/bin/env python3
"""
Extracteur de documents Word vers Markdown
- Extrait les chapitres présents dans la table des matières
- Conserve les titres, tableaux et bullet points
- Ignore les images, en-têtes et pieds de page
- Log les différences entre TDM et document
"""

import os
import re
import argparse
from pathlib import Path
from typing import List, Dict, Optional, Tuple, Set, Set
from dataclasses import dataclass, field
import logging

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class Chapter:
    """Représente un chapitre extrait"""
    title: str
    level: int
    content: List[str]


@dataclass
class ExtractionReport:
    """Rapport d'extraction avec les différences TDM/Document"""
    toc_entries: List[Tuple[str, int]] = field(default_factory=list)
    doc_headings: List[Tuple[str, int]] = field(default_factory=list)
    matched_chapters: List[str] = field(default_factory=list)
    in_toc_not_in_doc: List[str] = field(default_factory=list)
    in_doc_not_in_toc: List[str] = field(default_factory=list)


class WordToMarkdownExtractor:
    """Extracteur de contenu Word vers Markdown basé sur la TDM"""
    
    # Styles de titre Word courants
    HEADING_STYLES = [
        'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5',
        'Titre 1', 'Titre 2', 'Titre 3', 'Titre 4', 'Titre 5',
        'Heading1', 'Heading2', 'Heading3', 'Heading4', 'Heading5',
        'Title', 'Titre', 'Subtitle', 'Sous-titre'
    ]
    
    # Styles de liste Word courants
    LIST_STYLES = [
        'List Paragraph', 'List Bullet', 'List Number',
        'Paragraphe de liste', 'Liste à puces', 'Liste numérotée',
        'ListParagraph', 'ListBullet', 'ListNumber',
        'Bullet', 'Numbered', 'puce', 'numéro'
    ]
    
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.toc_entries: List[Tuple[str, int]] = []  # (titre, niveau)
        self.doc_headings: List[Tuple[str, int]] = []  # Titres trouvés dans le doc
        self.chapters: List[Chapter] = []
        self.report: ExtractionReport = ExtractionReport()
        
    def extract_toc(self) -> List[Tuple[str, int]]:
        """Extrait les entrées de la table des matières"""
        toc_entries = []
        in_toc = False
        
        for para in self.doc.paragraphs:
            # Détecter le début de la TDM
            para_text = para.text.strip()
            
            # Vérifier si c'est un champ TOC dans le XML
            if para._element.xml and 'w:fldChar' in para._element.xml:
                if 'TOC' in para._element.xml:
                    in_toc = True
                    continue
            
            # Détecter les styles de TDM
            style_name = para.style.name if para.style else ""
            
            if 'TOC' in style_name or 'TM ' in style_name or 'toc' in style_name.lower():
                # Nettoyer le texte (enlever les numéros de page)
                clean_title = self._clean_toc_entry(para_text)
                if clean_title:
                    level = self._get_toc_level(style_name)
                    toc_entries.append((clean_title, level))
                    logger.debug(f"TDM trouvée: [{level}] {clean_title}")
        
        # Si pas de TDM trouvée via les styles, chercher les titres directement
        if not toc_entries:
            logger.info("Pas de TDM détectée, extraction des titres du document...")
            toc_entries = self._extract_headings_as_toc()
        
        self.toc_entries = toc_entries
        logger.info(f"Entrées TDM extraites: {len(toc_entries)}")
        
        return toc_entries
    
    def _clean_toc_entry(self, text: str) -> str:
        """Nettoie une entrée de TDM (enlève numéros de page, tabs, etc.)"""
        # Enlever les numéros de page à la fin
        text = re.sub(r'\t+\d+\s*$', '', text)
        text = re.sub(r'\.{2,}\s*\d+\s*$', '', text)
        text = re.sub(r'\s+\d+\s*$', '', text)
        # Enlever les caractères de contrôle
        text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
        return text.strip()
    
    def _get_toc_level(self, style_name: str) -> int:
        """Détermine le niveau d'un style de TDM"""
        match = re.search(r'(\d+)', style_name)
        if match:
            return int(match.group(1))
        return 1
    
    def _extract_headings_as_toc(self) -> List[Tuple[str, int]]:
        """Extrait les titres du document comme TDM de secours"""
        headings = []
        
        for para in self.doc.paragraphs:
            style_name = para.style.name if para.style else ""
            
            for heading_style in self.HEADING_STYLES:
                if heading_style.lower() in style_name.lower():
                    level = self._get_heading_level(style_name)
                    title = para.text.strip()
                    if title:
                        headings.append((title, level))
                    break
        
        return headings
    
    def _get_heading_level(self, style_name: str) -> int:
        """Détermine le niveau d'un titre"""
        match = re.search(r'(\d+)', style_name)
        if match:
            return int(match.group(1))
        if 'subtitle' in style_name.lower() or 'sous-titre' in style_name.lower():
            return 2
        return 1
    
    def _normalize_title(self, title: str) -> str:
        """Normalise un titre pour comparaison"""
        # Enlever la numérotation au début
        normalized = re.sub(r'^[\d.]+\s*', '', title)
        # Enlever les espaces multiples
        normalized = re.sub(r'\s+', ' ', normalized)
        # Mettre en minuscules
        normalized = normalized.lower().strip()
        return normalized
    
    def _titles_match(self, title1: str, title2: str) -> bool:
        """Compare deux titres de manière flexible"""
        norm1 = self._normalize_title(title1)
        norm2 = self._normalize_title(title2)
        
        # Correspondance exacte
        if norm1 == norm2:
            return True
        
        # L'un contient l'autre
        if norm1 in norm2 or norm2 in norm1:
            return True
        
        # Correspondance partielle (80% des mots en commun)
        words1 = set(norm1.split())
        words2 = set(norm2.split())
        if words1 and words2:
            common = words1 & words2
            ratio = len(common) / max(len(words1), len(words2))
            if ratio >= 0.8:
                return True
        
        return False
    
    def extract_chapters(self) -> List[Chapter]:
        """Extrait le contenu des chapitres présents dans la TDM"""
        if not self.toc_entries:
            self.extract_toc()
        
        chapters = []
        current_chapter: Optional[Chapter] = None
        
        # Créer un set des titres normalisés de la TDM pour recherche rapide
        toc_titles_normalized = {self._normalize_title(t[0]): t for t in self.toc_entries}
        
        # Tracking pour le rapport
        matched_titles: Set[str] = set()
        doc_heading_titles: Set[str] = set()
        
        # Parcourir le document
        for element in self._iter_block_elements():
            if isinstance(element, Paragraph):
                para = element
                style_name = para.style.name if para.style else ""
                para_text = para.text.strip()
                
                # Ignorer les paragraphes vides
                if not para_text:
                    continue
                
                # Ignorer les entrées de TDM (on veut le vrai contenu)
                if 'TOC' in style_name or 'TM ' in style_name:
                    continue
                
                # Vérifier si c'est un titre
                is_heading = False
                heading_level = 0
                
                for heading_style in self.HEADING_STYLES:
                    if heading_style.lower() in style_name.lower():
                        is_heading = True
                        heading_level = self._get_heading_level(style_name)
                        break
                
                if is_heading:
                    # Ajouter aux titres du document
                    doc_heading_titles.add(self._normalize_title(para_text))
                    self.doc_headings.append((para_text, heading_level))
                    
                    # Vérifier si ce titre est dans la TDM
                    normalized = self._normalize_title(para_text)
                    in_toc = False
                    
                    for toc_title, toc_level in self.toc_entries:
                        if self._titles_match(para_text, toc_title):
                            in_toc = True
                            heading_level = toc_level
                            matched_titles.add(self._normalize_title(toc_title))
                            break
                    
                    if in_toc:
                        # Sauvegarder le chapitre précédent s'il existe
                        if current_chapter and current_chapter.content:
                            chapters.append(current_chapter)
                        
                        # Démarrer un nouveau chapitre
                        current_chapter = Chapter(
                            title=para_text,
                            level=heading_level,
                            content=[]
                        )
                        logger.debug(f"Chapitre trouvé: [{heading_level}] {para_text}")
                    elif current_chapter:
                        # Titre non dans TDM mais on est dans un chapitre
                        # L'ajouter comme sous-titre dans le contenu
                        md_heading = '#' * min(heading_level + 1, 6) + ' ' + para_text
                        current_chapter.content.append(md_heading)
                else:
                    # Contenu normal ou liste
                    if current_chapter:
                        # Vérifier si c'est une liste
                        list_item = self._parse_list_item(para)
                        if list_item:
                            current_chapter.content.append(list_item)
                        else:
                            current_chapter.content.append(para_text)
            
            elif isinstance(element, Table):
                # Convertir le tableau en Markdown
                if current_chapter:
                    table_md = self._table_to_markdown(element)
                    if table_md:
                        current_chapter.content.append(table_md)
        
        # Ajouter le dernier chapitre
        if current_chapter and current_chapter.content:
            chapters.append(current_chapter)
        
        self.chapters = chapters
        
        # Générer le rapport de comparaison
        self._generate_comparison_report(matched_titles, doc_heading_titles)
        
        logger.info(f"Chapitres extraits: {len(chapters)}")
        
        return chapters
    
    def _parse_list_item(self, para: Paragraph) -> Optional[str]:
        """Détecte et formate un élément de liste"""
        style_name = para.style.name if para.style else ""
        para_text = para.text.strip()
        
        if not para_text:
            return None
        
        # Vérifier le style
        is_list_style = any(ls.lower() in style_name.lower() for ls in self.LIST_STYLES)
        
        # Vérifier le XML pour les listes numérotées/à puces
        is_numbered = False
        is_bullet = False
        indent_level = 0
        
        # Chercher les propriétés de numérotation dans le XML
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                is_list_style = True
                # Niveau d'indentation
                ilvl = numPr.find(qn('w:ilvl'))
                if ilvl is not None:
                    indent_level = int(ilvl.get(qn('w:val'), 0))
                
                # Type de liste (numId peut aider à distinguer)
                numId = numPr.find(qn('w:numId'))
                if numId is not None:
                    num_id_val = numId.get(qn('w:val'))
                    # Généralement, les listes à puces ont des numId différents
                    # Mais c'est complexe, on se base sur le contenu
        
        # Détecter si le texte commence par un marqueur de liste
        bullet_patterns = [
            r'^[\u2022\u2023\u2043\u204C\u204D\u2219\u25AA\u25AB\u25CF\u25CB\u25D8\u25E6\u2605\u2606\u2610\u2611\u2612\u2713\u2714\u2715\u2716\u2717\u2718•●○◦‣⁃]\s*',
            r'^[-*+]\s+',
            r'^[–—]\s+',
        ]
        
        number_patterns = [
            r'^(\d+)[.)]\s+',
            r'^([a-zA-Z])[.)]\s+',
            r'^([ivxIVX]+)[.)]\s+',
            r'^(\d+\.\d+)[.)]\s*',
        ]
        
        for pattern in bullet_patterns:
            if re.match(pattern, para_text):
                is_bullet = True
                para_text = re.sub(pattern, '', para_text)
                break
        
        for pattern in number_patterns:
            match = re.match(pattern, para_text)
            if match:
                is_numbered = True
                para_text = re.sub(pattern, '', para_text)
                break
        
        # Formater en Markdown
        if is_list_style or is_bullet or is_numbered:
            indent = '  ' * indent_level
            if is_numbered:
                return f"{indent}1. {para_text}"
            else:
                return f"{indent}- {para_text}"
        
        return None
    
    def _generate_comparison_report(self, matched_titles: Set[str], doc_heading_titles: Set[str]):
        """Génère le rapport de comparaison TDM vs Document"""
        toc_normalized = {self._normalize_title(t[0]) for t in self.toc_entries}
        
        # Chapitres dans TDM mais pas dans le document
        in_toc_not_in_doc = []
        for toc_title, level in self.toc_entries:
            norm_title = self._normalize_title(toc_title)
            if norm_title not in matched_titles:
                in_toc_not_in_doc.append(toc_title)
        
        # Chapitres dans le document mais pas dans la TDM
        in_doc_not_in_toc = []
        for doc_title, level in self.doc_headings:
            norm_title = self._normalize_title(doc_title)
            found_in_toc = False
            for toc_title, _ in self.toc_entries:
                if self._titles_match(doc_title, toc_title):
                    found_in_toc = True
                    break
            if not found_in_toc:
                in_doc_not_in_toc.append(doc_title)
        
        # Mettre à jour le rapport
        self.report.toc_entries = self.toc_entries.copy()
        self.report.doc_headings = self.doc_headings.copy()
        self.report.matched_chapters = list(matched_titles)
        self.report.in_toc_not_in_doc = in_toc_not_in_doc
        self.report.in_doc_not_in_toc = in_doc_not_in_toc
        
        # Logger le rapport
        self._log_comparison_report()
    
    def _log_comparison_report(self):
        """Affiche le rapport de comparaison"""
        logger.info(f"\n{'='*60}")
        logger.info(f"RAPPORT DE COMPARAISON TDM vs DOCUMENT")
        logger.info(f"{'='*60}")
        
        logger.info(f"\n📋 Entrées dans la TDM: {len(self.report.toc_entries)}")
        logger.info(f"📄 Titres dans le document: {len(self.report.doc_headings)}")
        logger.info(f"✅ Chapitres correspondants: {len(self.report.matched_chapters)}")
        
        if self.report.in_toc_not_in_doc:
            logger.warning(f"\n⚠️  DANS TDM MAIS PAS DANS LE DOCUMENT ({len(self.report.in_toc_not_in_doc)}):")
            for title in self.report.in_toc_not_in_doc:
                logger.warning(f"   ❌ {title}")
        else:
            logger.info(f"\n✅ Tous les chapitres de la TDM sont présents dans le document")
        
        if self.report.in_doc_not_in_toc:
            logger.warning(f"\n⚠️  DANS DOCUMENT MAIS PAS DANS LA TDM ({len(self.report.in_doc_not_in_toc)}):")
            for title in self.report.in_doc_not_in_toc:
                logger.warning(f"   ❌ {title}")
        else:
            logger.info(f"\n✅ Tous les titres du document sont dans la TDM")
        
        logger.info(f"\n{'='*60}\n")
    
    def _iter_block_elements(self):
        """Itère sur les éléments du document (paragraphes et tableaux) dans l'ordre"""
        body = self.doc._body._body
        
        for child in body:
            if child.tag == qn('w:p'):
                yield Paragraph(child, self.doc._body)
            elif child.tag == qn('w:tbl'):
                yield Table(child, self.doc._body)
    
    def _table_to_markdown(self, table: Table) -> str:
        """Convertit un tableau Word en Markdown"""
        rows = []
        
        for row in table.rows:
            cells = []
            for cell in row.cells:
                # Extraire le texte de la cellule
                cell_text = cell.text.strip()
                # Échapper les pipes
                cell_text = cell_text.replace('|', '\\|')
                # Remplacer les retours à la ligne
                cell_text = cell_text.replace('\n', ' ')
                cells.append(cell_text)
            rows.append(cells)
        
        if not rows:
            return ""
        
        # Construire le Markdown
        md_lines = []
        
        # Première ligne (en-tête)
        md_lines.append('| ' + ' | '.join(rows[0]) + ' |')
        
        # Séparateur
        md_lines.append('| ' + ' | '.join(['---'] * len(rows[0])) + ' |')
        
        # Lignes de données
        for row in rows[1:]:
            # Gérer les lignes avec moins de colonnes
            while len(row) < len(rows[0]):
                row.append('')
            md_lines.append('| ' + ' | '.join(row[:len(rows[0])]) + ' |')
        
        return '\n'.join(md_lines)
    
    def to_markdown(self) -> str:
        """Convertit les chapitres extraits en Markdown"""
        if not self.chapters:
            self.extract_chapters()
        
        md_parts = []
        
        # Titre du document (nom du fichier)
        doc_name = Path(self.docx_path).stem
        md_parts.append(f"# {doc_name}\n")
        
        for chapter in self.chapters:
            # Titre du chapitre
            heading = '#' * min(chapter.level + 1, 6) + ' ' + chapter.title
            md_parts.append(f"\n{heading}\n")
            
            # Contenu
            for item in chapter.content:
                md_parts.append(item)
                md_parts.append("")  # Ligne vide entre les éléments
        
        return '\n'.join(md_parts)
    
    def save_markdown(self, output_path: Optional[str] = None) -> str:
        """Sauvegarde le Markdown dans un fichier"""
        if output_path is None:
            output_path = Path(self.docx_path).with_suffix('.md')
        
        md_content = self.to_markdown()
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        logger.info(f"Markdown sauvegardé: {output_path}")
        return str(output_path)
    
    def save_report(self, output_path: Optional[str] = None) -> str:
        """Sauvegarde le rapport de comparaison dans un fichier"""
        if output_path is None:
            output_path = Path(self.docx_path).with_suffix('.report.txt')
        
        lines = []
        lines.append("=" * 60)
        lines.append("RAPPORT DE COMPARAISON TDM vs DOCUMENT")
        lines.append(f"Fichier: {self.docx_path}")
        lines.append("=" * 60)
        lines.append("")
        
        lines.append(f"📋 Entrées dans la TDM: {len(self.report.toc_entries)}")
        lines.append(f"📄 Titres dans le document: {len(self.report.doc_headings)}")
        lines.append(f"✅ Chapitres correspondants: {len(self.report.matched_chapters)}")
        lines.append("")
        
        if self.report.in_toc_not_in_doc:
            lines.append(f"⚠️  DANS TDM MAIS PAS DANS LE DOCUMENT ({len(self.report.in_toc_not_in_doc)}):")
            for title in self.report.in_toc_not_in_doc:
                lines.append(f"   ❌ {title}")
            lines.append("")
        
        if self.report.in_doc_not_in_toc:
            lines.append(f"⚠️  DANS DOCUMENT MAIS PAS DANS LA TDM ({len(self.report.in_doc_not_in_toc)}):")
            for title in self.report.in_doc_not_in_toc:
                lines.append(f"   ❌ {title}")
            lines.append("")
        
        lines.append("-" * 60)
        lines.append("DÉTAIL TDM:")
        for title, level in self.report.toc_entries:
            lines.append(f"  [Niveau {level}] {title}")
        
        lines.append("")
        lines.append("-" * 60)
        lines.append("DÉTAIL TITRES DOCUMENT:")
        for title, level in self.report.doc_headings:
            lines.append(f"  [Niveau {level}] {title}")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        logger.info(f"Rapport sauvegardé: {output_path}")
        return str(output_path)


def process_directory(input_dir: str, output_dir: Optional[str] = None, 
                     recursive: bool = False, save_reports: bool = True) -> List[str]:
    """Traite tous les fichiers Word d'un répertoire"""
    input_path = Path(input_dir)
    
    if output_dir:
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
    else:
        output_path = input_path
    
    # Trouver les fichiers Word
    if recursive:
        docx_files = list(input_path.rglob('*.docx'))
    else:
        docx_files = list(input_path.glob('*.docx'))
    
    # Exclure les fichiers temporaires
    docx_files = [f for f in docx_files if not f.name.startswith('~$')]
    
    logger.info(f"Fichiers Word trouvés: {len(docx_files)}")
    
    output_files = []
    errors = []
    total_in_toc_not_doc = 0
    total_in_doc_not_toc = 0
    
    for docx_file in docx_files:
        try:
            logger.info(f"\nTraitement: {docx_file.name}")
            
            # Déterminer le chemin de sortie
            if output_dir:
                rel_path = docx_file.relative_to(input_path)
                md_file = output_path / rel_path.with_suffix('.md')
                report_file = output_path / rel_path.with_suffix('.report.txt')
                md_file.parent.mkdir(parents=True, exist_ok=True)
            else:
                md_file = docx_file.with_suffix('.md')
                report_file = docx_file.with_suffix('.report.txt')
            
            # Extraire et sauvegarder
            extractor = WordToMarkdownExtractor(str(docx_file))
            extractor.extract_toc()
            extractor.extract_chapters()
            saved_path = extractor.save_markdown(str(md_file))
            
            # Sauvegarder le rapport si demandé
            if save_reports:
                extractor.save_report(str(report_file))
            
            output_files.append(saved_path)
            
            # Accumuler les stats
            total_in_toc_not_doc += len(extractor.report.in_toc_not_in_doc)
            total_in_doc_not_toc += len(extractor.report.in_doc_not_in_toc)
            
            logger.info(f"  ✅ {len(extractor.chapters)} chapitres extraits")
            
        except Exception as e:
            logger.error(f"  ❌ Erreur: {e}")
            errors.append((str(docx_file), str(e)))
    
    # Résumé
    logger.info(f"\n{'='*60}")
    logger.info(f"RÉSUMÉ GLOBAL")
    logger.info(f"{'='*60}")
    logger.info(f"  - Fichiers traités avec succès: {len(output_files)}")
    logger.info(f"  - Fichiers en erreur: {len(errors)}")
    logger.info(f"  - Total chapitres TDM manquants dans docs: {total_in_toc_not_doc}")
    logger.info(f"  - Total titres docs manquants dans TDM: {total_in_doc_not_toc}")
    
    if errors:
        logger.info("\nFichiers en erreur:")
        for file, error in errors:
            logger.info(f"  - {file}: {error}")
    
    return output_files


def main():
    parser = argparse.ArgumentParser(
        description="Extrait le contenu des chapitres Word vers Markdown"
    )
    parser.add_argument(
        'input',
        help="Fichier Word ou répertoire à traiter"
    )
    parser.add_argument(
        '-o', '--output',
        help="Fichier ou répertoire de sortie (optionnel)"
    )
    parser.add_argument(
        '-r', '--recursive',
        action='store_true',
        help="Traiter les sous-répertoires"
    )
    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help="Afficher les détails de traitement"
    )
    parser.add_argument(
        '--no-report',
        action='store_true',
        help="Ne pas générer les fichiers de rapport"
    )
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    input_path = Path(args.input)
    
    if input_path.is_file():
        # Traiter un seul fichier
        if not input_path.suffix.lower() == '.docx':
            logger.error("Le fichier doit être au format .docx")
            return
        
        extractor = WordToMarkdownExtractor(str(input_path))
        
        logger.info(f"Extraction de: {input_path.name}")
        
        # Extraire la TDM
        toc = extractor.extract_toc()
        if toc:
            logger.info(f"\nTable des matières ({len(toc)} entrées):")
            for title, level in toc[:10]:
                logger.info(f"  {'  ' * (level-1)}[{level}] {title}")
            if len(toc) > 10:
                logger.info(f"  ... et {len(toc) - 10} autres entrées")
        
        # Extraire les chapitres
        chapters = extractor.extract_chapters()
        logger.info(f"\nChapitres extraits: {len(chapters)}")
        
        # Sauvegarder le markdown
        output_path = args.output if args.output else None
        saved = extractor.save_markdown(output_path)
        logger.info(f"\n✅ Markdown sauvegardé: {saved}")
        
        # Sauvegarder le rapport
        if not args.no_report:
            report_path = Path(saved).with_suffix('.report.txt')
            extractor.save_report(str(report_path))
        
    elif input_path.is_dir():
        # Traiter un répertoire
        process_directory(
            str(input_path),
            args.output,
            args.recursive,
            save_reports=not args.no_report
        )
    else:
        logger.error(f"Chemin invalide: {args.input}")


if __name__ == "__main__":
    main()
